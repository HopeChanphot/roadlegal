from __future__ import annotations

import json
import textwrap
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
OUT_PATH = OUT_DIR / "RoadLegal_Project_Technical_Documentation.docx"


SOURCE_FILES = [
    ".github/workflows/pages.yml",
    ".gitignore",
    "Dockerfile",
    "Procfile",
    "requirements.txt",
    "render.yaml",
    "railway.json",
    "roadlegal/__init__.py",
    "roadlegal/paths.py",
    "roadlegal/text.py",
    "roadlegal/llm_runtime.py",
    "roadlegal/challan.py",
    "roadlegal/geo.py",
    "roadlegal/game_content.py",
    "roadlegal/rag.py",
    "roadlegal/server.py",
    "scripts/download_sources.py",
    "scripts/build_index.py",
    "scripts/export_static_demo.py",
    "scripts/create_project_word_doc.py",
    "tests/test_roadlegal.py",
    "web/index.html",
    "web/styles.css",
    "web/app.js",
    "data/seed/source_manifest.json",
    "data/seed/fine_schedule.json",
    "data/seed/passages.json",
]

GENERATED_FILES = [
    "data/processed/passages.json",
    "web/static-data.json",
    "data/raw/downloads/*.html",
    "data/raw/downloads/*.pdf",
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)
                set_cell_margins(row.cells[idx])
                row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x1F, 0x29, 0x36)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Calibri"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        styles[name].font.bold = True
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    styles["Heading 1"].paragraph_format.space_before = Pt(18)
    styles["Heading 1"].paragraph_format.space_after = Pt(10)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    styles["Heading 2"].paragraph_format.space_before = Pt(14)
    styles["Heading 2"].paragraph_format.space_after = Pt(7)
    styles["Heading 3"].font.size = Pt(12)
    styles["Heading 3"].font.color.rgb = RGBColor(0x1F, 0x4D, 0x78)
    styles["Heading 3"].paragraph_format.space_before = Pt(10)
    styles["Heading 3"].paragraph_format.space_after = Pt(5)

    code = styles.add_style("CodeLine", 1)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code.font.size = Pt(7.2)
    code.font.color.rgb = RGBColor(0x21, 0x26, 0x2D)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.line_spacing = 1.0

    small = styles.add_style("SmallNote", 1)
    small.font.name = "Calibri"
    small._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    small.font.size = Pt(9)
    small.font.color.rgb = RGBColor(0x55, 0x5F, 0x6F)
    small.paragraph_format.space_after = Pt(4)


def add_footer(doc: Document) -> None:
    for section in doc.sections:
        para = section.footer.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = para.add_run("RoadLegal Technical Documentation")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x55, 0x5F, 0x6F)


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RoadLegal")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Development, Technical, Support, and Source Code Documentation")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    meta = doc.add_table(rows=5, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(meta, [1.75, 4.55])
    rows = [
        ("Project", "RoadLegal - offline-first BIMSTEC traffic-law and road-safety chatbot"),
        ("Repository", "https://github.com/HopeChanphot/roadlegal"),
        ("Local app", "http://127.0.0.1:8000/"),
        ("Public target", "https://hopechanphot.github.io/roadlegal/"),
        ("Generated", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")),
    ]
    for idx, (label, value) in enumerate(rows):
        meta.cell(idx, 0).text = label
        meta.cell(idx, 1).text = value
        set_cell_shading(meta.cell(idx, 0), "E8EEF5")
        for p in meta.cell(idx, 0).paragraphs:
            for r in p.runs:
                r.bold = True
    doc.add_paragraph()
    p = doc.add_paragraph(style="SmallNote")
    p.add_run(
        "This document is prepared for sharing project development details, technical architecture, package assumptions, "
        "deployment options, support information, and source-code listings. It is informational and not legal advice."
    )
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    doc.add_heading("Table of Contents", level=1)
    items = [
        "Executive Summary",
        "Problem Context and Product Idea",
        "Scope, Assumptions, and Constraints",
        "Software Packages and Tooling",
        "Architecture and Runtime Modes",
        "Data, Legal Review, and RAG Pipeline",
        "AI/LLM Runtime Strategy",
        "Frontend and User Experience Design",
        "API Reference",
        "Deployment and Hosting",
        "Testing and Verification",
        "Privacy, Security, and Safety",
        "Support and Maintenance",
        "Appendix A: Repository File Inventory",
        "Appendix B: Generated Data Summary",
        "Appendix C: Source Code Listings",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_page_break()


def add_callout(doc: Document, label: str, text: str, fill: str = "F4F6F9") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [6.3])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    para = cell.paragraphs[0]
    run = para.add_run(f"{label}: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    para.add_run(text)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, widths)
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for idx, header in enumerate(headers):
        cell = header_row.cells[idx]
        cell.text = header
        set_cell_shading(cell, "E8EEF5")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    set_table_width(table, widths)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_overview_sections(doc: Document) -> None:
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "RoadLegal is an offline-first AI road-safety and traffic-law assistant for BIMSTEC countries. It combines a "
        "local knowledge base, retrieval-augmented generation, structured challan calculation, country-aware UI, "
        "gamified learning, and public hosting paths. The application can run locally with a Python backend, online as a "
        "GitHub Pages static demo, or as a full backend service on cloud platforms such as Render or Railway."
    )
    add_callout(
        doc,
        "Important status",
        "The current MVP prioritizes source-backed legal information and cautious review flags. It does not claim to be legal advice.",
    )

    doc.add_heading("Problem Context and Product Idea", level=1)
    doc.add_paragraph(
        "Road traffic injuries are a major public-safety problem across South and Southeast Asia. The RoadLegal idea is "
        "to make traffic laws, penalties, safety explanations, and local enforcement information understandable through "
        "a conversational interface. The app also uses quizzes and scenarios because safety behavior improves when users "
        "engage with consequences, detection risk, and memorable practice situations."
    )
    add_bullets(
        doc,
        [
            "Primary users: drivers, riders, passengers, tourists, students, fleet operators, and road-safety educators.",
            "Core value: local legal guidance with citations and uncertainty labels.",
            "Hackathon differentiator: offline-first RAG and cross-country BIMSTEC switching.",
            "Behavior-change layer: quizzes, scenario games, score tracking, and safety coaching.",
        ],
    )

    doc.add_heading("Scope, Assumptions, and Constraints", level=1)
    add_table(
        doc,
        ["Area", "Current assumption or constraint"],
        [
            ["Legal advice", "The app is informational only and must not be presented as legal advice."],
            ["Fine accuracy", "Exact amounts are shown only where the starter record is source-backed; otherwise the app uses needs_review."],
            ["Offline use", "Static mode and local backend do not require paid cloud AI APIs."],
            ["Model runtime", "Generative answers require a local GGUF model compatible with llama.cpp."],
            ["GitHub Pages", "Static demo uses browser-side RAG over web/static-data.json."],
            ["Generated data", "Large generated JSON files are summarized in this document; source code is included in appendices."],
            ["Personal data", "No accounts are implemented; feedback is local in backend mode or browser localStorage in static mode."],
        ],
        [1.55, 4.75],
    )

    doc.add_heading("Software Packages and Tooling", level=1)
    add_table(
        doc,
        ["Tool or package", "Purpose", "Required for app?"],
        [
            ["Python 3", "Backend server, RAG engine, data scripts, tests.", "Yes for backend"],
            ["Python standard library", "HTTP server, JSON, paths, subprocess, text handling.", "Yes"],
            ["pypdf", "Optional PDF text extraction in scripts/build_index.py.", "Only for PDF ingestion"],
            ["llama.cpp", "Optional local GGUF model inference through llama-cli.", "Optional"],
            ["Git", "Version control and GitHub deployment.", "Yes for publishing"],
            ["GitHub Pages", "Static public demo hosting.", "Optional hosting"],
            ["Docker", "Portable backend deployment.", "Optional hosting"],
            ["Render/Railway", "Cloud backend deployment targets.", "Optional hosting"],
            ["Browser localStorage", "Static-mode score and feedback persistence.", "Yes for static UX"],
        ],
        [1.55, 3.6, 1.15],
    )

    doc.add_heading("Architecture and Runtime Modes", level=1)
    doc.add_paragraph("RoadLegal has three runtime modes:")
    add_numbered(
        doc,
        [
            "Backend mode: Python serves the API and web assets.",
            "Static mode: GitHub Pages serves web assets and browser-side fallback handles chat, calculator, and quiz.",
            "Generative mode: backend mode plus local llama.cpp GGUF generation when a model is available.",
        ],
    )
    add_table(
        doc,
        ["Layer", "Main files", "Responsibility"],
        [
            ["Frontend", "web/index.html, web/styles.css, web/app.js", "Chat UI, country switcher, quiz, calculator, static fallback."],
            ["API server", "roadlegal/server.py", "HTTP routes, static file serving, JSON responses."],
            ["RAG", "roadlegal/rag.py, roadlegal/text.py", "Lexical retrieval, citations, answer construction."],
            ["Calculator", "roadlegal/challan.py, data/seed/fine_schedule.json", "Fine estimates by offence and vehicle class."],
            ["Game", "roadlegal/game_content.py", "Quiz and country scenario content."],
            ["Geofence", "roadlegal/geo.py", "Starter BIMSTEC country matching."],
            ["Model", "roadlegal/llm_runtime.py", "Optional llama.cpp local generation."],
            ["Data scripts", "scripts/*.py", "Download, build index, export static demo."],
        ],
        [1.2, 2.45, 2.65],
    )

    doc.add_heading("Data, Legal Review, and RAG Pipeline", level=1)
    doc.add_paragraph(
        "The RAG corpus starts from curated seed passages and can be expanded with downloaded official or public legal "
        "documents. The processed index currently contains 658 passages."
    )
    add_numbered(
        doc,
        [
            "Add source URLs to data/seed/source_manifest.json.",
            "Run scripts/download_sources.py to fetch documents when network is available.",
            "Run scripts/build_index.py to extract and chunk text.",
            "Review legal meaning and update data/seed/fine_schedule.json.",
            "Run scripts/export_static_demo.py so GitHub Pages has updated packaged data.",
        ],
    )
    add_callout(
        doc,
        "Legal review policy",
        "When exact fine schedules are not verified from a current official source, records must stay marked as needs_review.",
        "FFF6DB",
    )

    doc.add_heading("AI/LLM Runtime Strategy", level=1)
    doc.add_paragraph(
        "RoadLegal is AI-enabled through RAG first. The model layer is optional. This avoids a fragile dependency on paid "
        "cloud APIs and keeps the demo usable offline. The local machine has llama-cli and llama-server, plus cached "
        "Llama 3.2 3B safetensors, but no GGUF model was found. Therefore the current backend mode is extractive-rag."
    )
    add_bullets(
        doc,
        [
            "To enable generation, place a compatible .gguf model in models/.",
            "Alternatively set ROADLEGAL_GGUF_MODEL to an absolute GGUF path.",
            "The backend will switch to generative-rag when llama-cli and GGUF are both available.",
            "Static GitHub Pages mode uses browser-side static-rag and cannot run llama.cpp.",
        ],
    )

    doc.add_heading("Frontend and User Experience Design", level=1)
    doc.add_paragraph(
        "The interface is intentionally practical rather than decorative. The first screen is the actual working tool: "
        "country selector, runtime status, chat, calculator, quiz, directory, and feedback."
    )
    add_table(
        doc,
        ["UI element", "Design reason"],
        [
            ["Country / law area menu", "Makes jurisdiction switching explicit and changes all dependent content."],
            ["Status strip", "Shows RAG mode, passage count, and model status for demo transparency."],
            ["Chat panel", "Primary conversational workflow."],
            ["Country profile card", "Explains current coverage and local assumptions."],
            ["Challan calculator", "Structured legal lookup independent from chat."],
            ["Quiz card", "Gamified retention and behavior-change layer."],
            ["Directory", "Practical enforcement and emergency links."],
            ["Feedback", "Legal data improvement loop."],
        ],
        [1.75, 4.55],
    )

    doc.add_heading("API Reference", level=1)
    add_table(
        doc,
        ["Endpoint", "Purpose"],
        [
            ["GET /api/health", "Runtime, model, index, and jurisdiction status."],
            ["GET /api/jurisdictions", "Country/law-area menu records."],
            ["GET /api/offences", "Offence options for selected jurisdiction."],
            ["POST /api/chat", "Grounded answer with citations and optional fine estimate."],
            ["POST /api/calculate-challan", "Structured fine estimate."],
            ["GET /api/geofence", "Country-level geofence lookup."],
            ["GET /api/quiz", "Country-aware quiz/scenario questions."],
            ["POST /api/feedback", "Local feedback logging."],
        ],
        [2.0, 4.3],
    )

    doc.add_heading("Deployment and Hosting", level=1)
    add_bullets(
        doc,
        [
            "GitHub repository: https://github.com/HopeChanphot/roadlegal",
            "Public static app target: https://hopechanphot.github.io/roadlegal/",
            "GitHub Pages branch: gh-pages, folder root.",
            "Full backend can be deployed with Docker, Render, or Railway.",
            "The backend reads the PORT environment variable for cloud platforms.",
        ],
    )
    add_callout(
        doc,
        "GitHub Pages setup",
        "If the public URL returns 404, enable Settings -> Pages -> Deploy from a branch -> gh-pages -> / (root).",
        "FFF6DB",
    )

    doc.add_heading("Testing and Verification", level=1)
    doc.add_paragraph("Current test command:")
    add_code_block(doc, "python -m unittest discover -s tests", include_heading=False)
    add_bullets(
        doc,
        [
            "India overspeeding fine test.",
            "Bangladesh needs_review caution test.",
            "Bangladesh geofence test.",
            "RAG citation test.",
            "Thailand expanded law and game content test.",
        ],
    )

    doc.add_heading("Privacy, Security, and Safety", level=1)
    add_bullets(
        doc,
        [
            "No paid cloud AI API is used by default.",
            "Backend feedback is appended to data/feedback.log.",
            "Static feedback is stored in browser localStorage only.",
            "No user accounts are implemented.",
            "Production deployment should add privacy notices, log redaction, and source-review signoff.",
        ],
    )

    doc.add_heading("Support and Maintenance", level=1)
    add_table(
        doc,
        ["Support topic", "Recommended response"],
        [
            ["Wrong fine", "Request current official source, update fine_schedule.json, rebuild static data, add test."],
            ["Missing country law", "Add source manifest record, seed passages, fine schedule records, and review flag."],
            ["GitHub Pages 404", "Enable Pages branch gh-pages root and wait several minutes."],
            ["Model not generative", "Add GGUF model or set ROADLEGAL_GGUF_MODEL."],
            ["Static app stale", "Run scripts/export_static_demo.py and push gh-pages branch."],
            ["Legal challenge", "Point to citation and review status; do not claim legal advice."],
        ],
        [1.6, 4.7],
    )

    doc.add_heading("Development Roadmap", level=1)
    add_bullets(
        doc,
        [
            "Add current official fine schedules for all BIMSTEC countries.",
            "Add municipal/state geofencing beyond rough country boxes.",
            "Add semantic retrieval with small embeddings.",
            "Add service-worker offline caching.",
            "Add voice input and text-to-speech.",
            "Add multilingual local translations.",
            "Add legal-review metadata and release versioning.",
            "Add authority/community content review workflow.",
        ],
    )


def add_inventory(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Appendix A: Repository File Inventory", level=1)
    tracked_like = []
    for path in sorted(ROOT.rglob("*")):
        rel = path.relative_to(ROOT).as_posix()
        if path.is_dir():
            continue
        if rel.startswith(".git/") or "__pycache__" in rel:
            continue
        if rel.startswith("data/raw/downloads/") and not rel.endswith(".gitkeep"):
            continue
        if rel.startswith("deliverables/"):
            continue
        tracked_like.append((rel, path.stat().st_size))
    rows = [[rel, f"{size:,} bytes"] for rel, size in tracked_like]
    add_table(doc, ["File", "Size"], rows, [4.8, 1.5])


def add_generated_summary(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Appendix B: Generated Data Summary", level=1)
    rows = []
    for rel in ["data/processed/passages.json", "web/static-data.json"]:
        path = ROOT / rel
        if path.exists():
            rows.append([rel, f"{path.stat().st_size:,} bytes", "Generated; summarized rather than printed in full."])
    rows.extend(
        [
            ["data/raw/downloads/*.html", "varies", "Ignored raw source downloads."],
            ["data/raw/downloads/*.pdf", "varies", "Ignored raw source downloads."],
            ["models/*.gguf", "large", "Ignored optional local model weights."],
        ]
    )
    add_table(doc, ["Generated artifact", "Size", "Reason"], rows, [2.45, 1.05, 2.8])
    doc.add_paragraph(
        "The complete app source code is included in Appendix C. Large generated RAG/static artifacts are excluded from "
        "the printed appendix because they are machine-generated from source data and would make the Word document less "
        "usable. They remain present in the repository."
    )


def add_code_block(doc: Document, code: str, include_heading: bool = True) -> None:
    if include_heading:
        doc.add_paragraph()
    for raw_line in code.splitlines() or [""]:
        line = raw_line.replace("\t", "    ")
        chunks = textwrap.wrap(
            line,
            width=112,
            replace_whitespace=False,
            drop_whitespace=False,
            break_long_words=True,
            break_on_hyphens=False,
        ) or [""]
        for idx, chunk in enumerate(chunks):
            prefix = "    " if idx else ""
            p = doc.add_paragraph(style="CodeLine")
            p.add_run(prefix + chunk)


def add_source_code_appendix(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Appendix C: Source Code Listings", level=1)
    doc.add_paragraph(
        "This appendix prints the handwritten source and configuration files needed to understand and rebuild the app. "
        "Generated data files are summarized in Appendix B."
    )
    for rel in SOURCE_FILES:
        path = ROOT / rel
        if not path.exists():
            continue
        doc.add_page_break()
        doc.add_heading(rel, level=2)
        try:
            content = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            content = path.read_text(encoding="utf-8", errors="replace")
        add_code_block(doc, content)


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_title(doc)
    add_toc(doc)
    add_overview_sections(doc)
    add_inventory(doc)
    add_generated_summary(doc)
    add_source_code_appendix(doc)
    add_footer(doc)
    doc.core_properties.title = "RoadLegal Development, Technical, Support, and Source Code Documentation"
    doc.core_properties.subject = "BIMSTEC road-safety AI chatbot technical documentation"
    doc.core_properties.author = "RoadLegal contributors"
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_doc()
